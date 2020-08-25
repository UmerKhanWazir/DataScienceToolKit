from flask import Flask, render_template, request, session, logging, url_for, redirect, flash
import requests
from sqlalchemy import create_engine
from sqlalchemy.orm import scoped_session, sessionmaker
from passlib.hash import sha256_crypt
from flask import send_from_directory
import PyPDF2
import pytesseract
from PIL import Image
import docx
from bs4 import BeautifulSoup
from urllib.request import urlopen
from ebooklib import epub
import ebooklib
import os
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'
import nltk
from mobi import mobi_uncompress
from werkzeug.utils import secure_filename
from flask_dance.contrib.github import make_github_blueprint, github
from authlib.integrations.flask_client import OAuth



engine = create_engine("mysql+pymysql://root:umeryusra@localhost/fyp")
db = scoped_session(sessionmaker(bind=engine))
app = Flask(__name__)

user_id = 1
page = ""
f_name = ""
UPLOAD_FOLDER = '/FYP/flask/data/files'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

github_blueprint = make_github_blueprint(client_id='e67539e6226d649008df',
                                              client_secret='5489afb74b14e586b89de1c1a5c3cebe0f819727')
app.register_blueprint(github_blueprint, url_prefix = '/github_login')

oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id="722544995990-4sfc9r2rjbr9o5bf93uv1mdhf8ji48t6.apps.googleusercontent.com",
    client_secret="KEbRAlj4ufF0N4pc99sOilcU",
    access_token_url='https://accounts.google.com/o/oauth2/token',
    access_token_params=None,
    authorize_url='https://accounts.google.com/o/oauth2/auth',
    authorize_params=None,
    api_base_url='https://www.googleapis.com/oauth2/v1/',
    userinfo_endpoint='https://openidconnect.googleapis.com/v1/userinfo',  # This is only needed if using openId to fetch user info
    client_kwargs={'scope': 'openid email profile'},
)

@app.route('/')
def front():
    return redirect(url_for('login'))

@app.route('/github-login')
def github_login():
    global user_id
    if not github.authorized:
        return redirect(url_for('github.login'))
    else:
        account_info = github.get('/user')
        if account_info.ok:
            account_info_json = account_info.json()
            user = account_info_json["id"]
            user1 = account_info_json["login"]
            usernamedata = db.execute("SELECT user_name FROM users WHERE user_name=:user_name", {"user_name": user}).fetchone()
            if usernamedata is None:
                db.execute("INSERT INTO users(name, user_name, profession, password) VALUES(:name, :user_name, :profession, :password)",
                                            {"name":user1, "user_name":user, "profession":"none", "password":"none"})
                db.commit()
                user_id = db.execute("SELECT id FROM users WHERE user_name=:user_name", {"user_name":user}).fetchone()
                return redirect(url_for("textsummarization"))
            else:
                user_id = db.execute("SELECT id FROM users WHERE user_name=:user_name", {"user_name":user}).fetchone()
                return redirect(url_for("textsummarization"))
            return redirect(url_for("textsummarization"))
    return '<h1>Request failed!</h1> '

@app.route('/google-login')
def google_login():
    google = oauth.create_client('google')
    redirect_uri = url_for('authorize', _external=True)
    return google.authorize_redirect(redirect_uri)

@app.route('/authorize')
def authorize():
    global user_id
    google = oauth.create_client('google')
    token = google.authorize_access_token()
    resp = google.get('userinfo')
    user_info = resp.json()
    user = user_info['email']
    user1 = user_info['name']
    usernamedata = db.execute("SELECT user_name FROM users WHERE user_name=:user_name", {"user_name": user}).fetchone()
    if usernamedata is None:
        db.execute("INSERT INTO users(name, user_name, profession, password) VALUES(:name, :user_name, :profession, :password)",
                                            {"name":user1, "user_name":user, "profession":"none", "password":"none"})
        db.commit()
        user_id = db.execute("SELECT id FROM users WHERE user_name=:user_name", {"user_name":user}).fetchone()
        return redirect(url_for("textsummarization"))
    else:
        user_id = db.execute("SELECT id FROM users WHERE user_name=:user_name", {"user_name":user}).fetchone()
        return redirect(url_for("textsummarization"))
    return redirect(url_for('textsummarization'))

@app.route('/signup', methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        name = request.form.get("fullname")
        user = request.form.get("username")
        profession = request.form.get("profession")
        password = request.form.get("password")
        confirm = request.form.get("confirmpassword")
        secure_password = sha256_crypt.encrypt(str(password))

        usernamedata = db.execute("SELECT user_name FROM users WHERE user_name=:user_name", {"user_name": user}).fetchone()

        if usernamedata is None:
            if(password == confirm):
                db.execute("INSERT INTO users(name, user_name, profession, password) VALUES(:name, :user_name, :profession, :password)",
                                            {"name":name, "user_name":user, "profession":profession, "password":secure_password})
                db.commit()
                flash("You are registered and can login now!","success")
                return redirect(url_for("login"))
            else:
                flash("Password does not match!","danger")
                return render_template("./Authentication/signup.html")
        else:
            flash("user name already taken!","danger")
            return render_template("./Authentication/signup.html")
    return render_template("./Authentication/signup.html")

@app.route('/login', methods=["GET", "POST"])
def login():
    if request.method == "POST":
        global user_id
        user_name = request.form.get("username1")
        password = request.form.get("password1")

        usernamedata = db.execute("SELECT user_name FROM users WHERE user_name=:user_name", {"user_name": user_name}).fetchone()
        passwordata = db.execute("SELECT password FROM users WHERE user_name=:user_name", {"user_name":user_name}).fetchone()
        user_id = db.execute("SELECT id FROM users WHERE user_name=:user_name", {"user_name":user_name}).fetchone()
        if usernamedata is None:
            flash("User not Found!","danger")
            return render_template("./Authentication/login.html")
        else:
            for passwor_data in passwordata:
                if sha256_crypt.verify(password,passwor_data):
                    flash("Congratulations! You have logged in successfully and can now use Data Science Toolkit.","success")
                    return redirect(url_for('textsummarization'))
                else:
                    flash("Incorrect password!","danger")
                    return render_template("./Authentication/login.html")
    return render_template("./Authentication/login.html")

@app.route('/text-summarization')
def textsummarization():
    global page
    name = db.execute("SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    profile = db.execute("SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    cover = db.execute("SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    return render_template("./sidebar-09/text_summarization.html", p=profile[0], c=cover[0], n=name[0], page = page )

@app.route('/upload-file', methods=["POST"])
def upload():
    global page
    global f_name
    files = request.files["fileToUpload"]
    filename = files.filename
    farray = filename.split(".")
    length = len(farray)
    ftype = length - 1
    filetype = farray[ftype]
    filetype = filetype.lower()
    if filetype == "png" or filetype == "jpg" or filetype == "jpeg" or filetype == "bmp" or filetype == "gif" or filetype == "tif": 
        files.save(os.path.join("/FYP/flask/data/files",filename))
        f_name = filename
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        img = Image.open(files)
        text = pytesseract.image_to_string(img)
        text = "".join([s for s in text.strip().splitlines(True) if s.strip()])
        page = text
        if request.form['submit'] == "Submit file for text-summarization":
            return redirect(url_for('textsummarization'))    
        elif request.form['submit'] == "save":
            filename = secure_filename(files.filename)
            name = f_name.split(".")
            name = name[0]+".txt"
            textfile = open(os.path.join("/FYP/flask/data/txt", name), "w")
            textfile.write(text)
            textfile.close()
            db.execute("INSERT INTO files(user_id, path, type) VALUES(:user_id, :path, :type)",
                                                    {"user_id":user_id[0], "path":f_name, "type":filetype})
            db.commit()
            flash('File saved!',"success")
            return redirect(url_for('repository'))
    
    elif filetype == "txt":
        files.save(os.path.join("/FYP/flask/data/files",filename))
        f_name = filename
        f = open(os.path.join("/FYP/flask/data/files", filename), "r")
        page = f.read()
        if request.form['submit'] == "Submit file for text-summarization":
            return redirect(url_for('textsummarization'))    
        elif request.form['submit'] == "save":
            filename = secure_filename(files.filename)
            name = f_name.split(".")
            name = name[0]+".txt"
            textfile = open(os.path.join("/FYP/flask/data/txt", name), "w")
            textfile.write(page)
            textfile.close()
            db.execute("INSERT INTO files(user_id, path, type) VALUES(:user_id, :path, :type)",
                                                    {"user_id":user_id[0], "path":f_name, "type":filetype})
            db.commit()
            flash('File saved!',"success")
            return redirect(url_for('repository'))
    elif filetype == "pdf":
        files.save(os.path.join("/FYP/flask/data/files",filename))
        f_name = filename
        pdfReader = PyPDF2.PdfFileReader(files)
        pages = pdfReader.numPages
        i=0
        text=""
        while i<pages:
            pageObject = pdfReader.getPage(i)
            text += pageObject.extractText()
            i = i+1
        page=text
        page = "".join([s for s in page.strip().splitlines(True) if s.strip()])
        if request.form['submit'] == "Submit file for text-summarization":
            return redirect(url_for('textsummarization'))    
        elif request.form['submit'] == "save":
            filename = secure_filename(files.filename)
            name = f_name.split(".")
            name = name[0]+".txt"
            textfile = open(os.path.join("/FYP/flask/data/txt", name), "w")
            textfile.write(text)
            textfile.close()
            db.execute("INSERT INTO files(user_id, path, type) VALUES(:user_id, :path, :type)",
                                                    {"user_id":user_id[0], "path":f_name, "type":filetype})
            db.commit()
            flash('File saved!',"success")
            return redirect(url_for('repository'))

    elif filetype == "docx" or filetype == "doc":
        files.save(os.path.join("/FYP/flask/data/files",filename))
        f_name = filename
        doc = docx.Document(files)
        paragraphs = len(doc.paragraphs)
        i = 0
        text = ""
        while ( i < paragraphs):
            text += doc.paragraphs[i].text
            i = i+1
            text += "\n"
        page = text
        page = "".join([s for s in page.strip().splitlines(True) if s.strip()])
        if request.form['submit'] == "Submit file for text-summarization":
            return redirect(url_for('textsummarization'))    
        elif request.form['submit'] == "save":
            filename = secure_filename(files.filename)
            name = f_name.split(".")
            name = name[0]+".txt"
            textfile = open(os.path.join("/FYP/flask/data/txt", name), "w")
            textfile.write(text)
            textfile.close()
            db.execute("INSERT INTO files(user_id, path, type) VALUES(:user_id, :path, :type)",
                                                    {"user_id":user_id[0], "path":f_name, "type":filetype})
            db.commit()
            flash('File saved!',"success")
            return redirect(url_for('repository'))

    elif filetype == "epub":
        files.save(os.path.join("/FYP/flask/data/files",filename))
        f_name = filename
        def epub2thtml(epub_path):
            book = epub.read_epub(epub_path)
            chapters = []
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    chapters.append(item.get_content())
            return chapters

        def chap2text(chap):
            output = ''
            soup = BeautifulSoup(chap, 'html.parser')
            text = soup.find_all(text=True)
            blacklist = [   '[document]',   'noscript', 'header',   'html', 'meta', 'head','input', 'script',   ]
            for t in text:
                if t.parent.name not in blacklist:
                    output += '{} '.format(t)
            return output

        def thtml2ttext(thtml):
            Output = []
            for html in thtml:
                text =  chap2text(html)
                Output.append(text)
            return Output

        def epub2text(epub_path):
            chapters = epub2thtml(epub_path)
            ttext = thtml2ttext(chapters)
            return ttext

        text = epub2text(files)
        length = len(text)
        i = 0
        text1 = ""
        while i < length:
            text1 += text[i]
            i = i+1
        text1 = "".join([s for s in text1.strip().splitlines(True) if s.strip()])
        page = text1
        if request.form['submit'] == "Submit file for text-summarization":
            return redirect(url_for('textsummarization'))    
        elif request.form['submit'] == "save":
            filename = secure_filename(files.filename)
            name = f_name.split(".")
            name = name[0]+".txt"
            textfile = open(os.path.join("/FYP/flask/data/txt", name), "w")
            textfile.write(page)
            textfile.close()
            db.execute("INSERT INTO files(user_id, path, type) VALUES(:user_id, :path, :type)",
                                                    {"user_id":user_id[0], "path":f_name, "type":filetype})
            db.commit()
            flash('File saved!',"success")
            return redirect(url_for('repository'))

    elif filetype == "mobi":
        book = mobi_uncompress(files)
        book.parse()
        text = ""
        for record in book:
            text += record
        page = text
        if request.form['submit'] == "Submit file for text-summarization":
            return redirect(url_for('textsummarization'))    
        elif request.form['submit'] == "save":
            filename = secure_filename(files.filename)
            name = filename.split(".")
            name = name[0]+".txt"
            files.save(os.path.join("/FYP/flask/data/files",filename))
            textfile = open(os.path.join("/FYP/flask/data/txt", name), "w")
            textfile.write(text)
            textfile.close()
            db.execute("INSERT INTO files(user_id, path, type) VALUES(:user_id, :path, :type)",
                                                    {"user_id":user_id[0], "path":filename, "type":filetype})
            db.commit()
            flash('File saved!',"success")
            return redirect(url_for('repository'))

@app.route('/web-scraping', methods=["POST"])
def web():
    global page
    url = request.form['url']
    html = urlopen(url).read()
    soup = BeautifulSoup(html)
    for script in soup(["script", "style"]):
        script.extract()  
    text = soup.get_text()
    lines = (line.strip() for line in text.splitlines())
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    text = '\n'.join(chunk for chunk in chunks if chunk)
    page = text
    return redirect(url_for('textsummarization'))

@app.route("/show-repository")
def showrepository():
    name = db.execute("SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    profile = db.execute("SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    cover = db.execute("SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    path = db.execute("SELECT path FROM files WHERE user_id =:user_id", {"user_id": user_id[0]}).fetchall()    
    return render_template("./sidebar-09/show_repository.html", n=name[0], p=profile[0], c=cover[0], path=path)

@app.route('/delete-file/<string:path>', methods= ['GET', 'POST'])
def deletefile(path):
    db.execute("DELETE FROM files WHERE path=:path", {"path":path})
    db.commit()
    return redirect(url_for('repository'))


@app.route('/upload-from-repository/<string:path>', methods= ['GET', 'POST'])
def uploadrepository(path):
    global page
    name = path.split(".")
    name = name[0]+".txt"
    f = open(os.path.join("/FYP/flask/data/txt", name), "r")
    page = f.read()
    return redirect(url_for("textsummarization"))

@app.route('/repository')
def repository():
    name = db.execute("SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    profile = db.execute("SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    cover = db.execute("SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    path = db.execute("SELECT path FROM files WHERE user_id =:user_id", {"user_id": user_id[0]}).fetchall()
    return render_template("./sidebar-09/repository.html", p=profile[0], n=name[0], c=cover[0], path = path)


@app.route('/view-file/<string:path>')
def view(path):
    return send_from_directory(app.config['UPLOAD_FOLDER'], path)

@app.route('/profile')
def user_info():
    name = db.execute("SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    profile = db.execute("SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    cover = db.execute("SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    return render_template("./sidebar-09/profile.html", p=profile[0], n=name[0], c=cover[0])

@app.route('/profile-picture',methods=["POST"])
def profile():
    files = request.files["profile"]
    filename = files.filename
    filename = secure_filename(files.filename)
    farray = filename.split(".")
    length = len(farray)
    ftype = length - 1
    filetype = farray[ftype]
    filetype = filetype.lower()
    profile = filename
    if filetype == "png" or filetype == "jpg" or filetype == "jpeg":
        output_size = (400,400)
        i = Image.open(files)
        i.thumbnail(output_size)
        i.save(os.path.join("/FYP/flask/static/sidebar/images/", filename))
        user1 = user_id[0]
        db.execute('UPDATE users SET profile=:filename WHERE id=:user1', {'filename': filename, 'user1': user1})
        db.commit()
        return redirect(url_for('user_info'))
    else:
        flash("Only png, jpg and jpeg files are accepted!","danger")
        return redirect(url_for('user_info'))

@app.route('/cover',methods=["POST"])
def cover():
    files = request.files["cover"]
    filename = files.filename
    filename = secure_filename(files.filename)
    farray = filename.split(".")
    length = len(farray)
    ftype = length - 1
    filetype = farray[ftype]
    filetype = filetype.lower()
    profile = filename
    if filetype == "png" or filetype == "jpg" or filetype == "jpeg":
        output_size = (400,400)
        i = Image.open(files)
        i.thumbnail(output_size)
        i.save(os.path.join("/FYP/flask/static/sidebar/images/", filename))
        user1 = user_id[0]
        db.execute('UPDATE users SET cover=:filename WHERE id=:user1', {'filename': filename, 'user1': user1})
        db.commit()
        return redirect(url_for('user_info'))
    else:
        flash("Only png, jpg and jpeg files are accepted!","danger")
        return redirect(url_for('user_info'))

@app.route('/users-search', methods=["GET", "POST"])
def users_search():
    user_name = request.form.get("search")
    user_name = user_name.split()
    words = len(user_name)
    i = words
    searched = []
    while (words>0):
        if not searched :
            if words == 1:
                searched = db.execute("SELECT * FROM users WHERE name LIKE '%"+user_name[0]+"%'  ").fetchall()
            elif words == 2:
                searched = db.execute("SELECT * FROM users WHERE name LIKE '%"+user_name[0]+"%' AND name LIKE '%"+user_name[1]+"%' ").fetchall()
            elif words == 3:
                searched = db.execute("SELECT * FROM users WHERE name LIKE '%"+user_name[0]+"%' AND name LIKE '%"+user_name[1]+"%' AND name LIKE '%"+user_name[2]+"%' ").fetchall()
            elif words == 4:
                searched = db.execute("SELECT * FROM users WHERE name LIKE '%"+user_name[0]+"%' AND name LIKE '%"+user_name[1]+"%' AND name LIKE '%"+user_name[2]+"%' AND name LIKE '%"+user_name[3]+"%' ").fetchall()
            elif words == 5:
                searched = db.execute("SELECT * FROM users WHERE name LIKE '%"+user_name[0]+"%' AND name LIKE '%"+user_name[1]+"%' AND name LIKE '%"+user_name[2]+"%' AND name LIKE '%"+user_name[3]+"%' AND name LIKE '%"+user_name[4]+"%'  ").fetchall()              
            else:
                searched = db.execute("SELECT * FROM users WHERE name LIKE '%"+user_name[0]+"%' AND name LIKE '%"+user_name[1]+"%' AND name LIKE '%"+user_name[2]+"%' AND name LIKE '%"+user_name[3]+"%' AND name LIKE '%"+user_name[4]+"%' AND name LIKE '%"+user_name[5]+"%' ").fetchall()      
        words = words-1
    if not searched:
        if i == 1:
            searched = db.execute("SELECT * FROM users WHERE name LIKE '%"+user_name[0]+"%'  ").fetchall()
        elif i == 2:
            searched = db.execute("SELECT * FROM users WHERE name LIKE '%"+user_name[0]+"%' OR name LIKE '%"+user_name[1]+"%' ").fetchall()
        elif i == 3:
            searched = db.execute("SELECT * FROM users WHERE name LIKE '%"+user_name[0]+"%' OR name LIKE '%"+user_name[1]+"%' OR name LIKE '%"+user_name[2]+"%' ").fetchall()
        elif i == 4:
            searched = db.execute("SELECT * FROM users WHERE name LIKE '%"+user_name[0]+"%' OR name LIKE '%"+user_name[1]+"%' OR name LIKE '%"+user_name[2]+"%' OR name LIKE '%"+user_name[3]+"%' ").fetchall()
        elif i == 5:
            searched = db.execute("SELECT * FROM users WHERE name LIKE '%"+user_name[0]+"%' OR name LIKE '%"+user_name[1]+"%' OR name LIKE '%"+user_name[2]+"%' OR name LIKE '%"+user_name[3]+"%' OR name LIKE '%"+user_name[4]+"%'  ").fetchall()              
        else:
            searched = db.execute("SELECT * FROM users WHERE name LIKE '%"+user_name[0]+"%' OR name LIKE '%"+user_name[1]+"%' OR name LIKE '%"+user_name[2]+"%' OR name LIKE '%"+user_name[3]+"%' OR name LIKE '%"+user_name[4]+"%' OR name LIKE '%"+user_name[5]+"%' ").fetchall()    

    name = db.execute("SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    profile = db.execute("SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    cover = db.execute("SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    return render_template("./sidebar-09/searched_users.html", n=name[0], p=profile[0], c=cover[0], searched=searched)

@app.route('/<string:user12>')
def user_12(user12):
    
    name = db.execute("SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    profile = db.execute("SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    cover = db.execute("SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    user = db.execute("SELECT * FROM users WHERE user_name=:user_name", {"user_name":user12}).fetchone()
    path = db.execute("SELECT path FROM files WHERE user_id =:user_id", {"user_id": user[0]}).fetchall()
    return render_template("./sidebar-09/searched_profile.html", n=name[0], p=profile[0], c=cover[0], u=user, path=path)
    
if __name__ == '__main__':
    app.secret_key="123456789umeryusra987654321"
    app.run(debug=True)